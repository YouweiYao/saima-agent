#!/usr/bin/env python3
"""赛马智能体 - 需求拆分"""
import sys
import json
import time
import requests
from concurrent.futures import ThreadPoolExecutor
from docx import Document

BASE_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"
API_KEY = "sk-3d96ade0c8fa40378a4560fdd43b067e"
MODEL = "qwen-plus"
MAX_WORKERS = 10

def call_llm(text: str) -> dict:
    """调用LLM"""
    prompt = """你是一位专业的标书需求分析师。

分类规则：
1. 功能性需求：技术功能、模块、能力
2. 非功能性需求：性能、安全、部署
3. 商务需求：采购、价格、合同
4. 运维/维保需求：7*24支持、问题解答
5. 验收需求：验收流程
6. 信创需求：国产化

需求点拆分规则：
- 一个需求点表达一个完整且独立的核心要求
- 不要做摘要，要拆分细化

输出格式：{"requirements":[{"requirement":"需求描述","category":"分类"}]}
只输出JSON！"""
    
    for attempt in range(3):
        try:
            response = requests.post(
                BASE_URL,
                headers={"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"},
                json={"model": MODEL, "messages": [{"role": "user", "content": f"{prompt}\n\n{text}"}]},
                timeout=60
            )
            if response.status_code == 200:
                data = response.json()
                content = data.get('choices', [{}])[0].get('message', {}).get('content', '')
                return json.loads(content)
        except Exception as e:
            print(f"调用错误: {e}")
    return None

def read_docx(path: str) -> str:
    """读取Word文档"""
    doc = Document(path)
    text = ""
    for p in doc.paragraphs:
        text += p.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
    return text

def split_text(text: str, chunk_size: int = 500) -> list:
    """分片"""
    chunks = []
    for i in range(0, len(text), chunk_size):
        chunks.append(text[i:i+chunk_size])
    return chunks

def process_chunk(chunk: str, idx: int, total: int) -> dict:
    """处理单个chunk"""
    print(f"处理 {idx+1}/{total}...")
    result = call_llm(chunk)
    if result:
        return {"source_text": chunk, "requirements": result.get("requirements", [])}
    return {"source_text": chunk, "requirements": []}

def main():
    if len(sys.argv) < 2:
        print("用法: python3 saima_main.py <word文件> [并发数]")
        return
    
    word_file = sys.argv[1]
    workers = int(sys.argv[2]) if len(sys.argv) > 2 else MAX_WORKERS
    
    print(f"读取文件: {word_file}")
    text = read_docx(word_file)
    
    print(f"分片...")
    chunks = split_text(text)
    print(f"共 {len(chunks)} 个片段")
    
    print(f"开始处理（并发{workers}）...")
    results = []
    with ThreadPoolExecutor(max_workers=workers) as executor:
        futures = [executor.submit(process_chunk, c, i, len(chunks)) for i, c in enumerate(chunks)]
        for f in futures:
            results.append(f.result())
    
    output = {"results": results}
    output_file = "/home/openclaw/niuniu/saima/output/需求拆分_result.json"
    with open(output_file, "w") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    
    print(f"完成！结果已保存到: {output_file}")

if __name__ == "__main__":
    main()
