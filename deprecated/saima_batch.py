#!/usr/bin/env python3
"""
赛马智能体 - 需求管理系统
优化版：批量处理减少模型调用
"""

import json
import sys
import os
from docx import Document
from typing import List, Dict

# ============ 步骤1: 读取Word文档 ============
def read_word_doc(file_path):
    """读取Word文档，返回文本列表"""
    doc = Document(file_path)
    texts = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text for cell in row.cells if cell.text.strip()])
            if row_text:
                texts.append(row_text)
    
    return texts


# ============ 步骤2: 文本分片 ============
def split_texts(texts, max_length=200):
    """将文本列表分片"""
    chunks = []
    current = ""
    
    for text in texts:
        text = text.strip()
        if not text:
            continue
        
        if len(text) > max_length:
            if current:
                chunks.append(current)
                current = ""
            sentences = text.replace("。", "。|").replace("\n", " ").split("|")
            for sent in sentences:
                sent = sent.strip()
                if not sent:
                    continue
                if len(current) + len(sent) + 1 <= max_length:
                    current += " " + sent if current else sent
                else:
                    if current:
                        chunks.append(current)
                    current = sent
        else:
            if len(current) + len(text) + 1 <= max_length:
                current += " " + text if current else text
            else:
                if current:
                    chunks.append(current)
                current = text
    
    if current:
        chunks.append(current)
    
    return chunks


# ============ 步骤3: 批量需求转化（优化版）============
def batch_process_requirements(chunks: List[str], batch_size: int = 5) -> List[Dict]:
    """
    批量处理需求，减少模型调用次数
    
    优化原理：多个分片一次输入，减少调用
    """
    results = []
    
    # 分批处理
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i+batch_size]
        
        # 构建批量prompt
        batch_text = "\n\n--- ---\n\n".join([
            f"【文本{i+1}】{chunk}" for i, chunk in enumerate(batch)
        ])
        
        # 这里调用模型（当前用模拟）
        # 实际使用时替换为真正的模型调用
        batch_results = simulate_llm_batch(batch)
        
        results.extend(batch_results)
        
        print(f"  已处理: {min(i+batch_size, len(chunks))}/{len(chunks)}")
    
    return results


def simulate_llm_batch(chunks: List[str]) -> List[Dict]:
    """
    模拟批量模型调用
    实际使用时替换为真正的模型调用
    """
    results = []
    
    for chunk in chunks:
        # 使用规则简单分类（可替换为真正模型调用）
        requirements = rule_based_classify(chunk)
        
        results.append({
            "source_text": chunk,
            "requirements": requirements
        })
    
    return results


def rule_based_classify(text: str) -> List[Dict]:
    """
    基于规则的需求分类
    作为批量处理的快速预处理
    """
    requirements = []
    
    # 关键词规则
    keywords = {
        "功能性需求": ["功能", "模块", "系统", "支持", "提供", "实现", "管理", "配置", "开发", "平台"],
        "安全需求": ["安全", "权限", "加密", "审计", "等保", "访问控制", "脱敏"],
        "性能需求": ["并发", "响应", "吞吐", "TPS", "QPS", "性能"],
        "部署需求": ["部署", "云化", "本地", "微服务", "容器", "K8s"],
        "维保需求": ["维保", "维护", "7×24", "容灾", "备份", "故障"],
        "信创需求": ["信创", "国产", "自主", "麒麟", "统信", "达梦", "飞腾", "鲲鹏", "龙芯"],
    }
    
    # 简单分类
    found_categories = set()
    for category, words in keywords.items():
        for word in words:
            if word in text:
                found_categories.add(category)
                break
    
    if not found_categories:
        found_categories.add("其他需求")
    
    # 生成需求条目
    for category in found_categories:
        requirements.append({
            "requirement": text[:100] + "..." if len(text) > 100 else text,
            "category": category
        })
    
    return requirements


def main():
    if len(sys.argv) < 2:
        print("用法: python3 saima_batch.py <word文件> [批次大小]")
        print("示例: python3 saima_batch.py 标书.docx 5")
        sys.exit(1)
    
    word_file = sys.argv[1]
    batch_size = int(sys.argv[2]) if len(sys.argv) > 2 else 5
    
    print("=" * 60)
    print("赛马智能体 - 需求管理系统（批量优化版）")
    print("=" * 60)
    print(f"输入文件: {word_file}")
    print(f"批次大小: {batch_size}")
    print()
    
    # 步骤1: 读取Word
    print("=== 步骤1: 读取Word文档 ===")
    texts = read_word_doc(word_file)
    print(f"✅ 段落数: {len(texts)}")
    print()
    
    # 步骤2: 分片
    print("=== 步骤2: 文本分片 ===")
    chunks = split_texts(texts, 200)
    print(f"✅ 分片数: {len(chunks)}")
    print(f"✅ 预计模型调用次数: {(len(chunks) + batch_size - 1) // batch_size} 次（原: {len(chunks)}次）")
    print()
    
    # 步骤3: 批量处理
    print("=== 步骤3: 批量需求转化 ===")
    results = batch_process_requirements(chunks, batch_size)
    print(f"✅ 处理完成")
    print()
    
    # 统计
    category_count = {}
    for r in results:
        for req in r.get("requirements", []):
            cat = req.get("category", "未知")
            category_count[cat] = category_count.get(cat, 0) + 1
    
    print("=== 分类统计 ===")
    for cat, count in sorted(category_count.items(), key=lambda x: -x[1]):
        print(f"  {cat}: {count}")
    
    # 保存结果
    output_file = "requirements_output.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print(f"\n结果已保存到: {output_file}")


if __name__ == "__main__":
    main()
