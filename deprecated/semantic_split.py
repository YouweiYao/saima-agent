#!/usr/bin/env python3
"""
语义分片 - 轻量级方案
使用 jieba 分词 + 简单相似度计算，无需大型 embedding 模型
"""

import re
import os
import sys
import json
import time
import argparse
import numpy as np
from pathlib import Path

# 依赖检查
try:
    import jieba
    HAS_JIEBA = True
except ImportError:
    HAS_JIEBA = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# 默认参数
DEFAULT_WINDOW_SIZE = 5
DEFAULT_STRIDE = 1
DEFAULT_MIN_CHUNK_CHARS = 500
DEFAULT_MAX_CHUNK_CHARS = 3000
DEFAULT_SIM_THRESHOLD = 0.15  # 相邻句子相似度阈值


def split_sentences(text: str) -> list:
    """按句子切分文本"""
    sentences = re.split(r'(?<=[。！？；\n])', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    return sentences


def split_long_sentence(sentence: str, max_len: int = 300) -> list:
    """超长句子按逗号再次切分"""
    if len(sentence) <= max_len:
        return [sentence]
    
    parts = re.split(r'(?<=[，、])', sentence)
    result = []
    current = ""
    for part in parts:
        if len(current) + len(part) <= max_len:
            current += part
        else:
            if current:
                result.append(current)
            current = part
    if current:
        result.append(current)
    return result if result else [sentence]


def load_text_from_docx(file_path: str) -> str:
    """从 Word 文档读取文本"""
    if not HAS_DOCX:
        import subprocess
        result = subprocess.run(
            ["pandoc", file_path, "-t", "plain"],
            capture_output=True, text=True
        )
        if result.returncode == 0:
            return result.stdout
        raise ImportError("python-docx not installed, and pandoc failed")
    
    doc = Document(file_path)
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    
    return "\n".join(paragraphs)


def load_text_from_file(file_path: str) -> str:
    """从文件读取文本"""
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    if suffix == ".docx":
        return load_text_from_docx(file_path)
    elif suffix == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def tokenize(text: str) -> set:
    """分词并返回词集合"""
    if not HAS_JIEBA:
        # 简单按字符分词
        return set(text)
    words = jieba.cut(text)
    return set(words)


def compute_word_similarity(text1: str, text2: str) -> float:
    """计算两个文本的词重叠相似度"""
    words1 = tokenize(text1)
    words2 = tokenize(text2)
    
    if not words1 or not words2:
        return 0.0
    
    intersection = len(words1 & words2)
    union = len(words1 | words2)
    
    return intersection / union if union > 0 else 0.0


def find_cut_points_by_similarity(sentences: list, threshold: float = 0.15) -> list:
    """根据相邻句子相似度找断点"""
    cut_points = []
    
    for i in range(len(sentences) - 1):
        sim = compute_word_similarity(sentences[i], sentences[i+1])
        
        # 如果相似度低于阈值，认为是语义断点
        if sim < threshold:
            cut_points.append(i)
    
    return cut_points


def generate_chunks(sentences: list, cut_points: list) -> list:
    """根据断点生成分片"""
    if not cut_points:
        return [{"text": "".join(sentences), "sentences": sentences}]
    
    chunks = []
    start = 0
    
    for cut in cut_points:
        chunk_text = "".join(sentences[start:cut+1])
        chunks.append({
            "text": chunk_text,
            "sentences": sentences[start:cut+1]
        })
        start = cut + 1
    
    # 最后一个 chunk
    if start < len(sentences):
        chunks.append({
            "text": "".join(sentences[start:]),
            "sentences": sentences[start:]
        })
    
    return chunks


def merge_small_chunks(chunks: list, min_chars: int = 500, max_chars: int = 3000) -> list:
    """合并太小的 chunks"""
    if not chunks:
        return chunks
    
    merged = []
    current = {"text": "", "sentences": []}
    
    for chunk in chunks:
        text = chunk["text"]
        
        # 如果当前 chunk 为空，直接使用
        if not current["text"]:
            current = chunk.copy()
            continue
        
        # 如果合并后不超过最大值，就合并
        if len(current["text"]) + len(text) <= max_chars:
            current["text"] += text
            current["sentences"].extend(chunk["sentences"])
        else:
            # 当前 chunk 满了，保存并新建
            if len(current["text"]) >= min_chars:
                merged.append(current)
            else:
                # 太小的合并到下一个
                if merged:
                    merged[-1]["text"] += current["text"]
                    merged[-1]["sentences"].extend(current["sentences"])
            current = chunk.copy()
    
    # 处理最后一个
    if current["text"]:
        if len(current["text"]) >= min_chars:
            merged.append(current)
        elif merged:
            merged[-1]["text"] += current["text"]
            merged[-1]["sentences"].extend(current["sentences"])
    
    return merged


def semantic_split(
    file_path: str,
    min_chars: int = DEFAULT_MIN_CHUNK_CHARS,
    max_chars: int = DEFAULT_MAX_CHUNK_CHARS,
    sim_threshold: float = DEFAULT_SIM_THRESHOLD,
    show_progress: bool = True
) -> list:
    """
    语义分片主函数
    """
    t0 = time.time()
    
    # 1. 读取文本
    if show_progress:
        print(f"[1/4] 读取文件: {file_path}")
    text = load_text_from_file(file_path)
    
    # 2. 句子切分
    if show_progress:
        print(f"[2/4] 句子切分...")
    sentences = split_sentences(text)
    
    # 超长句子再次切分
    all_sentences = []
    for sent in sentences:
        all_sentences.extend(split_long_sentence(sent))
    
    if show_progress:
        print(f"  -> 句子数: {len(all_sentences)}")
    
    if len(all_sentences) < 2:
        return [{
            "text": text,
            "sentences": all_sentences,
            "char_count": len(text)
        }]
    
    # 3. 找断点
    if show_progress:
        print(f"[3/4] 寻找语义断点 (阈值={sim_threshold})...")
    
    cut_points = find_cut_points_by_similarity(all_sentences, sim_threshold)
    
    if show_progress:
        print(f"  -> 找到 {len(cut_points)} 个断点")
    
    # 4. 生成分片
    if show_progress:
        print(f"[4/4] 生成分片...")
    
    chunks = generate_chunks(all_sentences, cut_points)
    chunks = merge_small_chunks(chunks, min_chars, max_chars)
    
    # 添加索引和字符数
    for i, chunk in enumerate(chunks):
        chunk["index"] = i
        chunk["char_count"] = len(chunk["text"])
    
    elapsed = time.time() - t0
    
    if show_progress:
        print(f"\n完成！耗时: {elapsed:.1f}秒")
        print(f"分片数: {len(chunks)}")
        total_chars = sum(c["char_count"] for c in chunks)
        print(f"总字符数: {total_chars}")
    
    return chunks


def main():
    parser = argparse.ArgumentParser(description="语义分片工具 (轻量级)")
    parser.add_argument("input", help="输入文件 (.txt, .docx)")
    parser.add_argument("-o", "--output", help="输出 JSON 文件路径")
    parser.add_argument("--min-chars", type=int, default=DEFAULT_MIN_CHUNK_CHARS)
    parser.add_argument("--max-chars", type=int, default=DEFAULT_MAX_CHUNK_CHARS)
    parser.add_argument("-t", "--threshold", type=float, default=DEFAULT_SIM_THRESHOLD)
    
    args = parser.parse_args()
    
    # 执行分片
    chunks = semantic_split(
        args.input,
        min_chars=args.min_chars,
        max_chars=args.max_chars,
        sim_threshold=args.threshold
    )
    
    # 输出
    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            json.dump(chunks, f, ensure_ascii=False, indent=2)
        print(f"结果已保存到: {args.output}")
    else:
        print("\n分片结果:")
        for i, chunk in enumerate(chunks):
            preview = chunk["text"][:100] + "..." if len(chunk["text"]) > 100 else chunk["text"]
            print(f"\n--- Chunk {i+1} ({chunk['char_count']} 字) ---")
            print(preview)


if __name__ == "__main__":
    main()
